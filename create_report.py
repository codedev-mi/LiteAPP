import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_background(cell, color_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)

def create_report():
    doc = Document()
    
    # Page setup
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # Styles Setup
    style_normal = doc.styles['Normal']
    font_normal = style_normal.font
    font_normal.name = 'Segoe UI'
    font_normal.size = Pt(11)
    font_normal.color.rgb = RGBColor(0x2D, 0x34, 0x36) # Dark grey
    
    # Colors
    color_primary = RGBColor(0x00, 0xB8, 0x94) # Bhusawal Mint Green
    color_dark = RGBColor(0x2D, 0x34, 0x36)
    
    # Document Header / Title
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run("BHUSAWAL BASKET - DELIVERABLES REPORT")
    title_run.font.bold = True
    title_run.font.size = Pt(22)
    title_run.font.color.rgb = color_primary
    
    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle_p.add_run("Project Status, Razorpay & UPI Payment Setup Guide, and Portability Solution\nTarget Deployment: 5 Days")
    sub_run.font.italic = True
    sub_run.font.size = Pt(12)
    sub_run.font.color.rgb = RGBColor(0x63, 0x6E, 0x72)
    
    doc.add_paragraph("\n")
    
    # 1. Executive Summary & Progress Chart
    h1 = doc.add_heading(level=1)
    run = h1.add_run("1. Module Completion Report")
    run.font.color.rgb = color_primary
    run.font.size = Pt(16)
    
    p = doc.add_paragraph()
    p.add_run("Below is the percentage value of completion for each application module within the Bhusawal Basket quick-commerce ecosystem. The codebase is heavily optimized, containing fully responsive designs, modular components, dynamic variant selectors, and error-resilient client-server routing.")
    
    table = doc.add_table(rows=5, cols=3)
    table.style = 'Light Shading Accent 1'
    
    headers = ["Module Name", "Completion %", "Status Summary"]
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        set_cell_background(hdr_cells[i], "00B894")
        for run in hdr_cells[i].paragraphs[0].runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            
    module_data = [
        ("Backend Services (Node.js/Prisma/PostgreSQL)", "95%", "Auth normalization, database, live notifications, Razorpay order APIs, and auto-assignments are 100% complete. Twilio/WhatsApp alert triggers are stubbed and ready for vendor API credentials."),
        ("User Mobile App (Expo/React Native)", "90%", "Integrated connection settings, Zomato tracking, post-delivery star review popup, promo calculations, platform fees (0.75%), and uniform product grids."),
        ("Admin Dashboard (Expo Web/Native)", "90%", "Clean tabbed navigation, variants creator, inventory tracking alerts, and status handlers. Pending final bundling checks."),
        ("Delivery Partner App (Expo/React Native)", "92%", "KYC uploading flow, online status, assignment alerts, and live location updates.")
    ]
    
    for row_idx, (name, val, summary) in enumerate(module_data, start=1):
        row_cells = table.rows[row_idx].cells
        row_cells[0].text = name
        row_cells[1].text = val
        row_cells[2].text = summary
        
        # Color the percentage cell text green
        for run in row_cells[1].paragraphs[0].runs:
            run.font.bold = True
            run.font.color.rgb = color_primary
            
    doc.add_paragraph("\n")
    
    # 2. Razorpay & UPI Payment Setup Guide
    h2 = doc.add_heading(level=1)
    run = h2.add_run("2. Razorpay & UPI Payment Setup Guide")
    run.font.color.rgb = color_primary
    run.font.size = Pt(16)
    
    doc.add_paragraph("To transition the application from a development environment to a live customer facing setup, the following payment integration steps are mandatory:")
    
    h2_sub1 = doc.add_heading(level=2)
    run = h2_sub1.add_run("A. Razorpay API Key Generation")
    run.font.color.rgb = color_dark
    run.font.size = Pt(13)
    
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Sign in to your ").font.bold = False
    p.add_run("Razorpay Dashboard").font.bold = True
    p.add_run(" (razorpay.com) and verify your KYC details.")
    
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Switch to ")
    p.add_run("Live Mode").font.bold = True
    p.add_run(" in the left side-panel.")
    
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Navigate to ").font.bold = False
    p.add_run("Settings -> API Keys -> Generate Key").font.bold = True
    p.add_run(".")
    
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Save the generated ")
    p.add_run("Key ID").font.bold = True
    p.add_run(" (starts with ")
    p.add_run("rzp_live_").font.italic = True
    p.add_run(") and ")
    p.add_run("Key Secret").font.bold = True
    p.add_run(" immediately.")
    
    h2_sub2 = doc.add_heading(level=2)
    run = h2_sub2.add_run("B. Environment Configuration")
    run.font.color.rgb = color_dark
    run.font.size = Pt(13)
    
    p = doc.add_paragraph()
    p.add_run("Update the following environment variables across the projects:")
    
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Backend (.env): ").font.bold = True
    p.add_run("Set ")
    p.add_run("RAZORPAY_KEY_ID=rzp_live_YOUR_KEY").font.italic = True
    p.add_run(" and ")
    p.add_run("RAZORPAY_KEY_SECRET=YOUR_SECRET").font.italic = True
    
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("User App (.env / eas.json): ").font.bold = True
    p.add_run("Set ")
    p.add_run("EXPO_PUBLIC_RAZORPAY_KEY_ID=rzp_live_YOUR_KEY").font.italic = True
    
    h2_sub3 = doc.add_heading(level=2)
    run = h2_sub3.add_run("C. Secure Payment Webhook Verification")
    run.font.color.rgb = color_dark
    run.font.size = Pt(13)
    
    p = doc.add_paragraph()
    p.add_run("In order to confirm payments safely (even if the user closes the app mid-transaction), you must configure a Webhook in the Razorpay Dashboard:")
    
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Webhook URL: ").font.bold = True
    p.add_run("https://your-domain.com/api/payment/webhook").font.italic = True
    
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Events to Subscribe: ").font.bold = True
    p.add_run("order.paid").font.italic = True
    p.add_run(", ")
    p.add_run("payment.failed").font.italic = True
    
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Secret: ").font.bold = True
    p.add_run("Define a random string and add it as ")
    p.add_run("RAZORPAY_WEBHOOK_SECRET").font.bold = True
    p.add_run(" in your backend environment variables.")
    
    doc.add_paragraph("\n")
    
    # 3. Portability Solution (Connection settings screen)
    h3 = doc.add_heading(level=1)
    run = h3.add_run("3. Application Portability & Multi-PC Solution")
    run.font.color.rgb = color_primary
    run.font.size = Pt(16)
    
    p = doc.add_paragraph()
    p.add_run("The primary reason why the Admin and Delivery applications previously failed to connect on different PCs was hardcoded IP addresses. When a PC switches Wi-Fi networks, its local IP changes, making the server unreachable by the mobile apps running on physical devices.")
    
    p_sol = doc.add_paragraph()
    p_sol.add_run("The Permanent Solution Implemented:").font.bold = True
    
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Dynamic URL Resolution Stack: ").font.bold = True
    p.add_run("On launch, the applications read from a priority stack starting with saved custom URLs in AsyncStorage, falling back to auto-generated host address URLs, and then localhost.")
    
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("ConnectionSettingsScreen: ").font.bold = True
    p.add_run("If the backend server is unreachable for 4 seconds, the app halts and automatically displays a sleek, dark-themed Connection Settings panel. It tells you exactly what is wrong, runs real-time latency tests, and offers auto-detected IP address suggestions so you can fix it with a single tap on any computer.")
    
    doc.add_paragraph("\n")
    
    # 4. Post-Delivery feedback (Zomato-style flow)
    h4 = doc.add_heading(level=1)
    run = h4.add_run("4. Zomato-style Order Flow")
    run.font.color.rgb = color_primary
    run.font.size = Pt(16)
    
    p = doc.add_paragraph()
    p.add_run("We have updated the checkout lifecycle to mimic Zomato's delivery satisfaction loop:")
    
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Instant Tracking: ").font.bold = True
    p.add_run("When the user places an order, the success popup navigates them immediately to the Live Tracking screen (DeliveryTrackingScreen.js) for that specific order ID.")
    
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Star Rating & Feedback Panel: ").font.bold = True
    p.add_run("As soon as the delivery status transitions to 'Delivered', the tracking screen dynamically transforms the delivery status timeline into a beautiful 'Rate & Review your items' card. Users can select 1-5 star ratings for each purchased product, write customized comments, and submit feedback straight to the database.")
    
    doc.add_paragraph("\n")
    
    # 5. Pending Items Checklist
    h5 = doc.add_heading(level=1)
    run = h5.add_run("5. Remaining Deployment Action Items")
    run.font.color.rgb = color_primary
    run.font.size = Pt(16)
    
    p = doc.add_paragraph()
    p.add_run("To proceed with deploying the application to friends or launching publicly in the next 5 days, please ensure the following tasks are completed:")
    
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("1. Compile and Sign Production Builds: ").font.bold = True
    p.add_run("Run ")
    p.add_run("eas build --platform android").font.italic = True
    p.add_run(" to generate the APK file to share in WhatsApp/Telegram groups so friends can install it directly.")
    
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("2. Production Server Hosting: ").font.bold = True
    p.add_run("Deploy the Node.js backend to Render/AWS/VPS and configure a permanent custom domain with an SSL certificate (HTTPS is required by the Razorpay SDK).")
    
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("3. WhatsApp Business API: ").font.bold = True
    p.add_run("Complete signup with a WhatsApp API provider (e.g. Twilio or Meta Business Manager) to replace placeholder console triggers with live SMS/WhatsApp stock alerts.")
    
    # Save document
    doc.save("c:/Users/Shree/Downloads/Telegram Desktop/LiteApp9/LiteApp_Project_Report.docx")
    print("Report generated successfully at c:/Users/Shree/Downloads/Telegram Desktop/LiteApp9/LiteApp_Project_Report.docx")

if __name__ == '__main__':
    create_report()
